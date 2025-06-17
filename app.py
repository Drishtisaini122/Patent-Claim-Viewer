from flask import Flask, render_template, request, send_file
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
import re
from dotenv import load_dotenv
import os

app = Flask(__name__)

load_dotenv()  # load variables from .env
API_KEY = os.getenv("API_KEY")

def set_landscape(section):
    section.orientation = 1  # landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    sectPr = section._sectPr
    pgSz = sectPr.xpath('./w:pgSz')[0]
    pgSz.set(qn('w:orient'), 'landscape')

def references_other_claims(text):
    return bool(re.search(r'\bclaim[s]?\s*\d+|\baccording to\b', text, re.IGNORECASE))

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        patent_number = request.form["patent_number"]

        search_url = "https://serpapi.com/search"
        params = {
            "engine": "google_patents",
            "q": patent_number,
            "api_key": API_KEY
        }
        search_response = requests.get(search_url, params=params)
        search_data = search_response.json()
        results = search_data.get("organic_results", [])

        if not results:
            return render_template("index.html", error="❌ No patent found")

        result = results[0]
        patent_link = result.get("patent_link")
        publication_number = result.get("publication_number")
        grant_date = result.get("grant_date", "N/A")
        patent_id = result.get("patent_id")

        details_url = "https://serpapi.com/search.json"
        details_params = {
            "engine": "google_patents_details",
            "patent_id": patent_id,
            "api_key": API_KEY
        }
        details_response = requests.get(details_url, params=details_params)
        details_data = details_response.json()

        abstract = details_data.get("abstract", "N/A")
        assignees = details_data.get("assignees", [])
        assignee = ", ".join(assignees) if assignees else "N/A"
        inventors_data = details_data.get("inventors", [])
        inventors = ", ".join([inv.get("name", "") for inv in inventors_data]) if inventors_data else "N/A"
        claims = details_data.get("claims", [])

        return render_template("index.html", patent={
            "link": patent_link,
            "publication_number": publication_number,
            "grant_date": grant_date,
            "abstract": abstract,
            "assignee": assignee,
            "inventors": inventors,
            "claims": claims,
            "title": details_data.get("title", "N/A"),
            "priority_date": details_data.get("priority_date", "N/A"),
            "filing_date": details_data.get("filing_date", "N/A")
        })

    return render_template("index.html")

@app.route("/download_docx", methods=["POST"])
def download_docx():
    claims = request.form.getlist("claims[]")
    publication_number = request.form.get("publication_number", "Patent")
    title = request.form.get("title", "")
    assignee = request.form.get("assignee", "")
    priority_date = request.form.get("priority_date", "")
    filing_date = request.form.get("filing_date", "")
    inventors = request.form.get("inventors", "")
    abstract = request.form.get("abstract", "")
    number_of_claims = str(len(claims))
    patent_link = request.form.get("patent_link", "")

    company_name = "Company Name"  # You can later replace this with request.form.get("company_name")

    doc = Document()
    section = doc.sections[0]
    set_landscape(section)

    # ===== Cover Page =====
    p0 = doc.add_paragraph()
    run0 = p0.add_run(company_name)
    run0.bold = True
    run0.font.size = Pt(28)
    p0.alignment = 1  # center

    p = doc.add_paragraph()
    run = p.add_run("Evidence of Use (EoU) Analysis")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = 1  # center

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Strictly Confidential")
    run2.bold = True
    run2.font.size = Pt(16)
    p2.alignment = 1  # center

    doc.add_page_break()

    # ===== Summary Info Page =====
    doc.add_heading(f"Patent Summary for {publication_number}", level=1)
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    summary_fields = [
        ("Publication Number:", publication_number),
        ("Title:", title),
        ("Assignee:", assignee),
        ("Priority Date:", priority_date),
        ("Filing Date:", filing_date),
        ("Inventors:", inventors),
        ("Abstract:", abstract),
        ("Number of Claims:", number_of_claims),
        ("Patent Link:", patent_link),
    ]
    for field, value in summary_fields:
        row_cells = info_table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = value

    doc.add_page_break()
    p1 = doc.add_paragraph()
    run1 = p1.add_run(f"Preliminary Analysis of Infringement of U.S. Patent No. {publication_number}")
    run1.bold = True
    run1.font.size = Pt(11)
    p1.alignment = 1  # center

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Strictly Confidential")
    run2.bold = True
    run2.font.size = Pt(11)
    p2.alignment = 1  # center

    doc.add_paragraph()  # spacing

    p3 = doc.add_paragraph()
    run3 = p3.add_run("About the Defendant: Company Name")
    run3.bold = True
    run3.underline = True
    run3.font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    run4 = p4.add_run("Source:")
    run4.bold = True

    # ===== Add Accused Instrumentality Page =====
    doc.add_page_break()
    p5 = doc.add_paragraph()
    run5 = p5.add_run("Accused Instrumentality: Product Name")
    run5.bold = True
    run5.underline = True
    run5.font.size = Pt(14)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p6 = doc.add_paragraph()
    run6 = p6.add_run("Source:")
    run6.font.size = Pt(11)
    run6.bold = True


    doc.add_page_break()

    # ===== Claim Chart =====
    doc.add_heading(f"Claim Chart for {publication_number}", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Claim Element'
    hdr_cells[1].text = 'Evidence'

    for row in table.rows:
        row.cells[0].width = Inches(4)
        row.cells[1].width = Inches(4)
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for claim in claims:
        lines = [line.strip() for line in claim.split('\n') if line.strip()]
        if not lines:
            continue

        claim_number_match = re.match(r'^(\d+)\.?\s*(.*)', lines[0])
        if claim_number_match:
            base_number = claim_number_match.group(1)
            claim_intro = claim_number_match.group(2).strip()
        else:
            base_number = "X"
            claim_intro = lines[0]

        row_cells = table.add_row().cells
        row_cells[0].text = f"{base_number} {claim_intro}"
        row_cells[1].text = ""
        row = table.rows[-1]
        row.height = Inches(0.4)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        sub_number = 1
        for line in lines[1:]:
            stripped_line = line.rstrip()
            if not re.search(r';\s*(and|or)\s*$', stripped_line) and not stripped_line.endswith('.') and not stripped_line.endswith(';'):
                line = stripped_line + ';'
            else:
                line = stripped_line

            label = f"{base_number}.{sub_number} {line}"
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = ""
            sub_number += 1

            row = table.rows[-1]
            row.height = Inches(0.4)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{publication_number}_claim_chart.docx"
    return send_file(buffer, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    app.run(debug=True)
